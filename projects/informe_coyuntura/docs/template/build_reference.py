"""
Construye el reference.docx institucional CIGOB partiendo del default de pandoc.

Paleta institucional CIGOB (extraída de cigob.org):
  - Azul principal:    #1D67CD
  - Azul oscuro:       #0F3D7A
  - Gris texto:        #333333
  - Gris secundario:   #5F6360

Uso:
    cd projects/informe_coyuntura/docs/template
    pandoc -o cigob_reference.docx --print-default-data-file reference.docx
    python build_reference.py
    # genera cigob_reference.docx con estilos CIGOB + logo en header
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def find_style(doc, name):
    """python-docx __getitem__ tiene bugs; iterar manualmente y devolver por nombre."""
    for s in doc.styles:
        if s.name == name:
            return s
    return None

HERE = Path(__file__).parent
SRC = HERE / "cigob_reference.docx"
LOGO = HERE / "cigob_logo.png"

# Paleta institucional CIGOB
AZUL_CIGOB     = RGBColor(0x1D, 0x67, 0xCD)
AZUL_OSCURO    = RGBColor(0x0F, 0x3D, 0x7A)
GRIS_TEXTO     = RGBColor(0x33, 0x33, 0x33)
GRIS_SECUNDARIO = RGBColor(0x5F, 0x63, 0x60)

FONT_TITULOS = "Calibri"
FONT_CUERPO  = "Calibri"
FONT_MONO    = "Consolas"


def apply_font_xml(rPr, name):
    """Forzar rFonts a nivel XML para que Word respete la fuente en todos los scripts."""
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def run_font(run, name, size=None, color=None, bold=None, italic=None):
    """Configura una fuente sobre un Run."""
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    # XML
    rPr = run._element.get_or_add_rPr()
    apply_font_xml(rPr, name)


def style_font(doc, style_name, font, size, color=None, bold=False, italic=False, space_before=None, space_after=None):
    """Configura completamente un estilo de párrafo por nombre."""
    st = find_style(doc, style_name)
    if st is None:
        return False
    f = st.font
    f.name = font
    f.size = Pt(size)
    if color is not None:
        f.color.rgb = color
    f.bold = bold
    f.italic = italic
    rPr = st.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        st.element.append(rPr)
    apply_font_xml(rPr, font)
    if (space_before is not None or space_after is not None):
        pPr = st.element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            st.element.append(pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        if space_before is not None:
            spacing.set(qn("w:before"), str(int(space_before * 20)))
        if space_after is not None:
            spacing.set(qn("w:after"), str(int(space_after * 20)))
    return True


def add_border_below(style_element, color_hex="1D67CD", size_pt=1):
    """Agrega línea inferior decorativa al estilo (para Heading 1)."""
    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_element.append(pPr)
    # Quitar bordes previos si los hay
    old = pPr.find(qn("w:pBdr"))
    if old is not None:
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt * 8))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def replace_header_with_logo(doc, logo_path):
    """Inserta logo CIGOB en el header de la primera sección."""
    section = doc.sections[0]
    header = section.header
    # Vaciar contenido previo
    for p in list(header.paragraphs):
        for r in list(p.runs):
            r.text = ""
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(str(logo_path), height=Cm(1.0))
    # Línea separadora azul debajo del header
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn("w:pBdr"))
    if old is not None:
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), "1D67CD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_footer(doc):
    """Footer institucional: 'Fundación CiGob · página X / Y'."""
    section = doc.sections[0]
    footer = section.footer
    for p in list(footer.paragraphs):
        for r in list(p.runs):
            r.text = ""
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Línea superior gris
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn("w:pBdr"))
    if old is not None:
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), "BFBFBF")
    pBdr.append(top)
    pPr.append(pBdr)
    # Texto
    run_left = p.add_run("Fundación CiGob   ·   página ")
    run_font(run_left, FONT_CUERPO, 8, GRIS_SECUNDARIO)
    # Campo PAGE
    run_page = p.add_run()
    run_font(run_page, FONT_CUERPO, 8, GRIS_SECUNDARIO)
    for elem in (
        ("w:fldChar", {"w:fldCharType": "begin"}),
        ("w:instrText", None, "PAGE"),
        ("w:fldChar", {"w:fldCharType": "end"}),
    ):
        if len(elem) == 3:
            tag, attrs, text = elem
            el = OxmlElement(tag)
            el.text = text
        else:
            tag, attrs = elem
            el = OxmlElement(tag)
            if attrs:
                for k, v in attrs.items():
                    el.set(qn(k), v)
        run_page._r.append(el)
    # Separador
    run_sep = p.add_run(" de ")
    run_font(run_sep, FONT_CUERPO, 8, GRIS_SECUNDARIO)
    # Campo NUMPAGES
    run_total = p.add_run()
    run_font(run_total, FONT_CUERPO, 8, GRIS_SECUNDARIO)
    for elem in (
        ("w:fldChar", {"w:fldCharType": "begin"}),
        ("w:instrText", None, "NUMPAGES"),
        ("w:fldChar", {"w:fldCharType": "end"}),
    ):
        if len(elem) == 3:
            tag, attrs, text = elem
            el = OxmlElement(tag)
            el.text = text
        else:
            tag, attrs = elem
            el = OxmlElement(tag)
            if attrs:
                for k, v in attrs.items():
                    el.set(qn(k), v)
        run_total._r.append(el)


def configure_table_style(doc):
    """Bordes finos color CIGOB para todas las tablas vía estilo 'Table'."""
    st = find_style(doc, "Table")
    if st is None:
        return
    tblPr = st.element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        st.element.insert(0, tblPr)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def main():
    if not SRC.exists():
        raise SystemExit(f"Falta {SRC}. Generar primero con: pandoc -o cigob_reference.docx --print-default-data-file reference.docx")

    doc = Document(str(SRC))

    # ── Estilos ────────────────────────────────────────────────────────────────
    applied = []
    def s(*args, **kw):
        if style_font(doc, *args, **kw):
            applied.append(args[0])

    # Título principal (# en md → "Title")
    s("Title",     FONT_TITULOS, 26, color=AZUL_OSCURO, bold=True,  space_before=0,  space_after=8)
    s("Subtitle",  FONT_TITULOS, 14, color=GRIS_SECUNDARIO, italic=True, space_before=0, space_after=18)
    # Headings
    s("Heading 1", FONT_TITULOS, 18, color=AZUL_CIGOB,  bold=True, space_before=24, space_after=8)
    s("Heading 2", FONT_TITULOS, 14, color=AZUL_CIGOB,  bold=True, space_before=18, space_after=6)
    s("Heading 3", FONT_TITULOS, 12, color=AZUL_OSCURO, bold=True, space_before=12, space_after=4)
    s("Heading 4", FONT_TITULOS, 11, color=GRIS_TEXTO,  bold=True, italic=True, space_before=10, space_after=4)
    s("Heading 5", FONT_TITULOS, 11, color=GRIS_TEXTO,  bold=True, space_before=8, space_after=3)
    s("Heading 6", FONT_TITULOS, 11, color=GRIS_SECUNDARIO, italic=True, space_before=8, space_after=3)
    # Cuerpo
    s("Normal",          FONT_CUERPO, 11, color=GRIS_TEXTO, space_before=0, space_after=6)
    s("Body Text",       FONT_CUERPO, 11, color=GRIS_TEXTO, space_before=0, space_after=6)
    s("First Paragraph", FONT_CUERPO, 11, color=GRIS_TEXTO, space_before=0, space_after=6)
    s("Compact",         FONT_CUERPO, 11, color=GRIS_TEXTO, space_before=0, space_after=2)
    # Autor / Fecha / Abstract
    s("Author", FONT_CUERPO, 11, color=GRIS_SECUNDARIO, italic=True, space_before=0, space_after=2)
    s("Date",   FONT_CUERPO, 11, color=GRIS_SECUNDARIO, italic=True, space_before=0, space_after=18)
    s("Abstract", FONT_CUERPO, 11, color=GRIS_TEXTO, italic=True, space_before=12, space_after=12)
    # Citas (Block Text es el de pandoc para blockquotes)
    s("Block Text", FONT_CUERPO, 11, color=GRIS_SECUNDARIO, italic=True, space_before=8, space_after=8)
    # Caption
    s("Caption",       FONT_CUERPO, 9, color=GRIS_SECUNDARIO, italic=True, space_before=2, space_after=8)
    s("Image Caption", FONT_CUERPO, 9, color=GRIS_SECUNDARIO, italic=True, space_before=2, space_after=8)
    s("Table Caption", FONT_CUERPO, 9, color=GRIS_SECUNDARIO, italic=True, space_before=2, space_after=8)
    # Footnotes
    s("Footnote Text", FONT_CUERPO, 9, color=GRIS_TEXTO, space_before=0, space_after=2)
    # Definition list
    s("Definition Term", FONT_TITULOS, 11, color=AZUL_OSCURO, bold=True, space_before=6, space_after=0)
    s("Definition",      FONT_CUERPO, 11, color=GRIS_TEXTO, space_before=0, space_after=6)

    # Heading 1 con línea decorativa azul
    h1 = find_style(doc, "Heading 1")
    if h1 is not None:
        add_border_below(h1.element, "1D67CD")

    # Verbatim Char (inline code de pandoc)
    vst = find_style(doc, "Verbatim Char")
    if vst is not None:
        vst.font.name = FONT_MONO
        vst.font.size = Pt(9)
        vst.font.color.rgb = AZUL_OSCURO
        rPr = vst.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            vst.element.append(rPr)
        apply_font_xml(rPr, FONT_MONO)

    hst = find_style(doc, "Hyperlink")
    if hst is not None:
        hst.font.color.rgb = AZUL_CIGOB

    # ── Tablas ────────────────────────────────────────────────────────────────
    configure_table_style(doc)

    # ── Header + footer ──────────────────────────────────────────────────────
    if LOGO.exists():
        replace_header_with_logo(doc, LOGO)
    configure_footer(doc)

    # ── Márgenes institucionales ─────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    out = HERE / "cigob_reference.docx"
    doc.save(str(out))
    print(f"OK {out.name} actualizado")
    print(f"   estilos aplicados ({len(applied)}): {', '.join(applied)}")


if __name__ == "__main__":
    main()
