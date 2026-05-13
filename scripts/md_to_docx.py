"""
md_to_docx.py — Conversor Markdown → Word para CIGOB
Uso:
  python scripts/md_to_docx.py archivo.md
  python scripts/md_to_docx.py archivo1.md archivo2.md --output-dir output/
  python scripts/md_to_docx.py docs/carpeta/              # convierte todos los .md
Requiere: python-docx  (pip install python-docx)
"""
import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

# ── Paleta CIGOB ──────────────────────────────────────────────────────────────
C_AZUL_OSCURO  = RGBColor(0x1F, 0x39, 0x64)   # H1
C_AZUL_MEDIO   = RGBColor(0x2E, 0x74, 0xB5)   # H2
C_AZUL_CLARO   = RGBColor(0x2F, 0x52, 0x96)   # H3
C_GRIS_TEXTO   = RGBColor(0x40, 0x40, 0x40)   # párrafos
C_GRIS_CODIGO  = RGBColor(0xF2, 0xF2, 0xF2)   # fondo bloque código
C_BORDE_TABLA  = RGBColor(0xBD, 0xD7, 0xEE)
FONT_BODY      = "Calibri"
FONT_CODE      = "Courier New"


# ── Helpers de estilo ─────────────────────────────────────────────────────────

def _set_font(run, name, size_pt, bold=False, italic=False, color=None):
    run.font.name       = name
    run.font.size       = Pt(size_pt)
    run.font.bold       = bold
    run.font.italic     = italic
    if color:
        run.font.color.rgb = color


def _shade_paragraph(para, hex_fill: str):
    """Aplica color de fondo (shading) a un párrafo."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   kwargs.get("val", "single"))
        el.set(qn("w:sz"),    kwargs.get("sz",  "4"))
        el.set(qn("w:color"), kwargs.get("color", "BDD7EE"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _apply_inline(run_text: str, para, font_size: float, color=None):
    """Parsea **bold**, *italic*, `code` dentro de un texto y agrega runs al párrafo."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(run_text):
        if m.start() > last:
            r = para.add_run(run_text[last:m.start()])
            _set_font(r, FONT_BODY, font_size, color=color or C_GRIS_TEXTO)
        if m.group(2):                    # **bold**
            r = para.add_run(m.group(2))
            _set_font(r, FONT_BODY, font_size, bold=True, color=color or C_GRIS_TEXTO)
        elif m.group(3):                  # *italic*
            r = para.add_run(m.group(3))
            _set_font(r, FONT_BODY, font_size, italic=True, color=color or C_GRIS_TEXTO)
        elif m.group(4):                  # `code`
            r = para.add_run(m.group(4))
            _set_font(r, FONT_CODE, font_size - 1, color=RGBColor(0xC0, 0x39, 0x2B))
        last = m.end()
    if last < len(run_text):
        r = para.add_run(run_text[last:])
        _set_font(r, FONT_BODY, font_size, color=color or C_GRIS_TEXTO)


# ── Document setup ────────────────────────────────────────────────────────────

def _build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)

    # Estilo Normal base
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(4)

    return doc


# ── Parseo de Markdown ────────────────────────────────────────────────────────

def _strip_frontmatter(lines: list[str]) -> list[str]:
    if not lines or lines[0].strip() != "---":
        return lines
    try:
        end = lines.index("---\n", 1)
        return lines[end + 1:]
    except ValueError:
        try:
            end = lines.index("---", 1)
            return lines[end + 1:]
        except ValueError:
            return lines


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and "|" in line.strip()[1:]


def _parse_table(block: list[str], doc: Document):
    rows = [r for r in block if not re.match(r"^\|[-| :]+\|$", r.strip())]
    if not rows:
        return
    col_count = max(len(r.strip().strip("|").split("|")) for r in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    for i, raw in enumerate(rows):
        cells_text = [c.strip() for c in raw.strip().strip("|").split("|")]
        while len(cells_text) < col_count:
            cells_text.append("")
        row = table.add_row()
        for j, ct in enumerate(cells_text[:col_count]):
            cell = row.cells[j]
            _set_cell_border(cell)
            para = cell.paragraphs[0]
            is_header = i == 0
            if is_header:
                cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2E74B5")
                cell._tc.tcPr.append(shd)
                _apply_inline(ct, para, 10, color=RGBColor(0xFF, 0xFF, 0xFF))
                for run in para.runs:
                    run.font.bold = True
            else:
                _apply_inline(ct, para, 10)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()


def convert_md_to_docx(md_path: Path, output_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
    lines = _strip_frontmatter(raw)

    doc = _build_document()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip("\n").rstrip()

        # ── Encabezados ───────────────────────────────────────────────────────
        m_h = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m_h:
            level   = len(m_h.group(1))
            text    = m_h.group(2)
            para    = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(8 if level == 1 else 5)
            para.paragraph_format.space_after  = Pt(4)
            if level == 1:
                para.paragraph_format.space_before = Pt(14)
                r = para.add_run(text)
                _set_font(r, FONT_BODY, 18, bold=True, color=C_AZUL_OSCURO)
                # Línea inferior decorativa
                pPr = para._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:color"), "1F3964")
                pBdr.append(bottom)
                pPr.append(pBdr)
            elif level == 2:
                r = para.add_run(text)
                _set_font(r, FONT_BODY, 14, bold=True, color=C_AZUL_MEDIO)
            elif level == 3:
                r = para.add_run(text)
                _set_font(r, FONT_BODY, 12, bold=True, color=C_AZUL_CLARO)
            else:
                r = para.add_run(text)
                _set_font(r, FONT_BODY, 11, bold=True, color=C_GRIS_TEXTO)
            i += 1
            continue

        # ── Bloque de código ──────────────────────────────────────────────────
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].rstrip("\n").startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # cierre ```
            for cl in code_lines:
                para = doc.add_paragraph()
                _shade_paragraph(para, "F2F2F2")
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)
                para.paragraph_format.left_indent  = Cm(0.5)
                r = para.add_run(cl)
                _set_font(r, FONT_CODE, 9, color=RGBColor(0x20, 0x20, 0x20))
            doc.add_paragraph()
            continue

        # ── Tabla ─────────────────────────────────────────────────────────────
        if _is_table_row(stripped):
            table_block = []
            while i < len(lines) and _is_table_row(lines[i].rstrip("\n").rstrip()):
                table_block.append(lines[i].rstrip("\n"))
                i += 1
            _parse_table(table_block, doc)
            continue

        # ── Lista con viñetas ─────────────────────────────────────────────────
        m_li = re.match(r"^(\s*)[-*]\s+(.*)", stripped)
        if m_li:
            indent = len(m_li.group(1)) // 2
            text   = m_li.group(2)
            para   = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(0.5 + indent * 0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            # Limpiar el run que agrega List Bullet automáticamente
            for run in para.runs:
                para._p.remove(run._r)
            _apply_inline(text, para, 11)
            i += 1
            continue

        # ── Lista numerada ────────────────────────────────────────────────────
        m_nl = re.match(r"^\d+\.\s+(.*)", stripped)
        if m_nl:
            text = m_nl.group(1)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent  = Cm(0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            for run in para.runs:
                para._p.remove(run._r)
            _apply_inline(text, para, 11)
            i += 1
            continue

        # ── Línea horizontal ──────────────────────────────────────────────────
        if re.match(r"^---+$", stripped) or re.match(r"^\*\*\*+$", stripped):
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "BDD7EE")
            pBdr.append(bottom)
            pPr.append(pBdr)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if stripped.startswith("> "):
            text = stripped[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Cm(1.0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            _shade_paragraph(para, "EBF3FB")
            _apply_inline(text, para, 10.5, color=RGBColor(0x2E, 0x74, 0xB5))
            i += 1
            continue

        # ── Línea vacía ───────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Párrafo normal ────────────────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(4)
        _apply_inline(stripped, para, 11)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[OK] {md_path} → {output_path}")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Convierte archivos Markdown a Word (.docx) con formato CIGOB")
    parser.add_argument("inputs", nargs="+", help="Archivos .md o directorios")
    parser.add_argument("--output-dir", "-o", default=None, help="Carpeta destino (default: misma carpeta que el .md)")
    args = parser.parse_args()

    md_files: list[Path] = []
    for inp in args.inputs:
        p = Path(inp)
        if p.is_dir():
            md_files.extend(sorted(p.glob("**/*.md")))
        elif p.is_file() and p.suffix == ".md":
            md_files.append(p)
        else:
            print(f"[SKIP] {inp} — no es .md ni directorio")

    if not md_files:
        print("Sin archivos .md para convertir.")
        sys.exit(1)

    for md in md_files:
        if args.output_dir:
            out = Path(args.output_dir) / (md.stem + ".docx")
        else:
            out = md.with_suffix(".docx")
        try:
            convert_md_to_docx(md, out)
        except Exception as e:
            print(f"[ERROR] {md}: {e}")


if __name__ == "__main__":
    main()
